import os
from PyPDF2 import PdfReader
from openai import OpenAI
from tqdm import tqdm
from docx import Document
import re

class PDFTextProcessor:
    def __init__(self, api_key_file, prompt_file):
        self.api_key = self._read_api_key(api_key_file)
        self.prompt = self._read_prompt_file(prompt_file)
        self.client = OpenAI(api_key=self.api_key)

    @staticmethod
    def _read_api_key(file_path):
        with open(file_path, "r") as file:
            return file.read().strip()

    @staticmethod
    def _read_prompt_file(file_path):
        with open(file_path, "r") as file:
            return file.read().strip()

    def extract_text_from_pdf(self, pdf_path):
        """Extract all text from the PDF file."""
        reader = PdfReader(pdf_path)
        text = ""
        for page in reader.pages:
            text += page.extract_text()
        return text

    def clean_text(self, text):
        """Clean up the extracted text by removing excessive newlines."""
        return re.sub(r'\n+', '\n', text).strip()

    def split_text_by_token_limit(self, text, max_tokens=30000):
        """Split the text into chunks respecting the token limit and completing sentences."""
        words = text.split()
        chunks = []
        current_chunk = []
        current_tokens = 0

        for word in words:
            word_tokens = len(word) // 4 + 1  # Rough estimate: 1 token ~ 4 characters
            current_chunk.append(word)
            current_tokens += word_tokens

            if current_tokens > max_tokens:
                # Join current chunk and find the last sentence boundary
                combined_chunk = " ".join(current_chunk)
                last_period_index = combined_chunk.rfind('.')
                if last_period_index == -1:
                    # If no period is found, split at max_tokens
                    chunks.append(combined_chunk)
                    current_chunk = []
                    current_tokens = 0
                else:
                    # Split at the last sentence boundary
                    chunks.append(combined_chunk[:last_period_index + 1])
                    remaining_text = combined_chunk[last_period_index + 1:].strip()
                    current_chunk = remaining_text.split()
                    current_tokens = sum(len(word) // 4 + 1 for word in current_chunk)

        # Append any remaining chunk
        if current_chunk:
            chunks.append(" ".join(current_chunk))

        return chunks

    def rewrite_text(self, text, ref_text):
        """Send a request to the OpenAI API to rewrite the text."""

        user_message = {
            'role': 'user',
            'content': self.prompt + ref_text + text
        }
        try:
            response = self.client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[user_message],
            )
            return response.choices[0].message.content.strip()
        except Exception as e:
            print(f"An error occurred: {e}")
            return "Error in processing the request."

    def save_chunks_to_single_doc(self, chunks, output_path):
        """Save all rewritten chunks to a single Word file."""
        doc = Document()
        for i, chunk in enumerate(chunks):
            doc.add_heading(f"Chunk {i + 1}", level=1)
            doc.add_paragraph(chunk)
        doc.save(output_path)

    def process_pdf(self, pdf_path, ref_path):
        """Process the PDF: extract, clean, split, rewrite, and save to Word."""
        print("Extracting text from PDF...")
        text = self.extract_text_from_pdf(pdf_path)


        # -- refrencing the new PDF --

        ref_text = self.extract_text_from_pdf(ref_path)

        # -- refrencing the new PDF --
        
        # print("Cleaning extracted text...")
        # text = self.clean_text(text)

        print("Splitting text into chunks...")
        chunks = self.split_text_by_token_limit(text, max_tokens=30000)

        rewritten_chunks = []

        print("Rewriting chunks using OpenAI API...")
        for i, chunk in enumerate(tqdm(chunks, desc="Rewriting chunks")):
            rewritten_chunk = self.rewrite_text(chunk, ref_text)
            rewritten_chunks.append(rewritten_chunk)

        output_path = "Rewritten_Document.docx"
        print(f"Saving all rewritten chunks to {output_path}...")
        self.save_chunks_to_single_doc(rewritten_chunks, output_path)

        print("Processing complete.")

# Main script
if __name__ == "__main__":
    api_key_file = "API_KEY.txt"
    prompt_file = "prompt.txt"
    pdf_path = "Report - Disclosure 1-12-24 to USMS.pdf"
    ref_path = "RULE 235 DEFAULT PROCEEDING AND DISPOSITIVE MOTION-1.pdf"

    processor = PDFTextProcessor(api_key_file, prompt_file)
    processor.process_pdf(pdf_path, ref_path)
